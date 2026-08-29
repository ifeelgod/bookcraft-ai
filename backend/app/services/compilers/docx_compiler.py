"""
DOCX Compiler implementation using python-docx.
Compiles DocumentAST into fully-styled, editable Microsoft Word (.docx) documents.
Preserves headings, styles, callouts, pullquotes, tables, interactive fields, and front matter.
"""
from __future__ import annotations
import asyncio
import logging
import re
from pathlib import Path
from typing import Optional, Tuple

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor

from app.models.document_ast import (
    CalloutBlock,
    Chapter,
    DocumentAST,
    Heading2Block,
    Heading3Block,
    HorizontalRuleBlock,
    ImageBlock,
    InteractiveFieldBlock,
    PageBreakBlock,
    ParagraphBlock,
    PullquoteBlock,
    TableBlock,
    TrimSize,
)
from app.services.compilers.base import BaseCompiler, sanitize_filename

logger = logging.getLogger("bookcraft.compiler.docx")

# Trim size mapping in inches: (width, height)
TRIM_DIMENSIONS_DOCX = {
    TrimSize.small: (5.5, 8.5),
    TrimSize.medium: (6.0, 9.0),
    TrimSize.large: (8.5, 11.0),
}

# Callout theme colors: (fill_hex, border_hex, text_color_hex)
CALLOUT_COLORS = {
    "info": ("EFF6FF", "1D4ED8", "1E40AF"),      # Blue
    "tip": ("F0FDF4", "166534", "14532D"),       # Green
    "warning": ("FFFBEB", "92400E", "78350F"),   # Amber
    "danger": ("FEF2F2", "991B1B", "7F1D1D"),    # Red
    "success": ("F0FDF4", "166534", "14532D"),   # Emerald
}

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _add_formatted_text(paragraph, text: str, default_bold: bool = False, default_italic: bool = False):
    """
    Parse inline markdown tokens (***bold-italic***, **bold**, *italic*, `code`)
    and append formatted Runs to the given docx paragraph.
    """
    if not text:
        return

    # Tokenizer pattern matching bold-italic, bold, italic, code spans
    pattern = r'(\*\*\*.*?\*\*\*|___.*?___|\*\*.*?\*\*|__.*?__|\*[^\*]+?\*|_[^_]+?_|`[^`]+?`)'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue

        if (part.startswith('***') and part.endswith('***') and len(part) >= 6) or \
           (part.startswith('___') and part.endswith('___') and len(part) >= 6):
            inner = part[3:-3]
            run = paragraph.add_run(inner)
            run.bold = True
            run.italic = True
        elif (part.startswith('**') and part.endswith('**') and len(part) >= 4) or \
             (part.startswith('__') and part.endswith('__') and len(part) >= 4):
            inner = part[2:-2]
            run = paragraph.add_run(inner)
            run.bold = True
            run.italic = default_italic
        elif (part.startswith('*') and part.endswith('*') and len(part) >= 2) or \
             (part.startswith('_') and part.endswith('_') and len(part) >= 2):
            inner = part[1:-1]
            run = paragraph.add_run(inner)
            run.bold = default_bold
            run.italic = True
        elif part.startswith('`') and part.endswith('`') and len(part) >= 2:
            inner = part[1:-1]
            run = paragraph.add_run(inner)
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.bold = default_bold
            run.italic = default_italic
        else:
            run = paragraph.add_run(part)
            run.bold = default_bold
            run.italic = default_italic


def _set_cell_background(cell, fill_hex: str):
    """Set the background shading of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tc_pr.append(shd)


def _set_cell_borders(
    cell,
    top: Optional[str] = None,
    bottom: Optional[str] = None,
    left: Optional[str] = None,
    right: Optional[str] = None,
    sz: str = "4",
):
    """Set custom borders on a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders_xml = f'<w:tcBorders {nsdecls("w")}>'
    borders_xml += f'<w:top w:val="{"single" if top else "none"}" w:sz="{sz}" w:space="0" w:color="{top or "auto"}"/>'
    borders_xml += f'<w:left w:val="{"single" if left else "none"}" w:sz="{sz}" w:space="0" w:color="{left or "auto"}"/>'
    borders_xml += f'<w:bottom w:val="{"single" if bottom else "none"}" w:sz="{sz}" w:space="0" w:color="{bottom or "auto"}"/>'
    borders_xml += f'<w:right w:val="{"single" if right else "none"}" w:sz="{sz}" w:space="0" w:color="{right or "auto"}"/>'
    borders_xml += '</w:tcBorders>'
    tc_pr.append(parse_xml(borders_xml))


def _set_cell_margins(cell, top: int = 140, bottom: int = 140, left: int = 200, right: int = 200):
    """Set cell internal padding in twips (1 pt = 20 twips)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tc_pr.append(tc_mar)


def _build_docx_document(ast: DocumentAST) -> Document:
    """Build a complete python-docx Document instance from DocumentAST."""
    doc = Document()
    cs = ast.compilation_settings

    # ── Page Geometry & Margins ───────────────────────────────────────────────
    section = doc.sections[0]
    trim_w, trim_h = TRIM_DIMENSIONS_DOCX.get(ast.metadata.trim_size, (6.0, 9.0))
    section.page_width = Inches(trim_w)
    section.page_height = Inches(trim_h)

    # Margins from compilation settings
    margins = cs.margins
    section.top_margin = Inches(margins.top if margins else 0.7)
    section.bottom_margin = Inches(margins.bottom if margins else 0.7)
    section.left_margin = Inches(margins.inner if margins else 0.75)
    section.right_margin = Inches(margins.outer if margins else 0.5)

    # Header / Footer settings
    section.different_first_page_header_footer = True
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if cs.header_footer.show_book_title:
        hrun = header_para.add_run(ast.metadata.title)
        hrun.font.size = Pt(8.5)
        hrun.font.italic = True
        hrun.font.color.rgb = RGBColor(128, 128, 128)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Default Styles ────────────────────────────────────────────────────────
    font_name = cs.font_family.value if hasattr(cs.font_family, 'value') else str(cs.font_family or 'Garamond')
    style_normal = doc.styles['Normal']
    style_normal.font.name = font_name
    style_normal.font.size = Pt(cs.font_size or 11)
    style_normal.font.color.rgb = RGBColor(30, 30, 30)
    style_normal.paragraph_format.line_spacing = cs.line_height or 1.35
    style_normal.paragraph_format.space_after = Pt(4)

    # ── Front Matter: Title Page ──────────────────────────────────────────────
    if ast.front_matter.title_page.enabled:
        # Vertical space
        for _ in range(3):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)

        # Title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_after = Pt(12)
        run_title = p_title.add_run(ast.metadata.title)
        run_title.font.name = font_name
        run_title.font.size = Pt(26)
        run_title.bold = True
        run_title.font.color.rgb = RGBColor(17, 24, 39)

        # Subtitle
        if ast.metadata.subtitle:
            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sub.paragraph_format.space_after = Pt(24)
            run_sub = p_sub.add_run(ast.metadata.subtitle)
            run_sub.font.name = font_name
            run_sub.font.size = Pt(15)
            run_sub.italic = True
            run_sub.font.color.rgb = RGBColor(75, 85, 99)

        # Author
        for _ in range(2):
            doc.add_paragraph()

        p_author = doc.add_paragraph()
        p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_author.paragraph_format.space_after = Pt(8)
        run_author = p_author.add_run(ast.metadata.author)
        run_author.font.name = font_name
        run_author.font.size = Pt(14)
        run_author.bold = True

        if ast.metadata.publisher:
            p_pub = doc.add_paragraph()
            p_pub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_pub.paragraph_format.space_after = Pt(4)
            run_pub = p_pub.add_run(ast.metadata.publisher)
            run_pub.font.size = Pt(10)
            run_pub.font.color.rgb = RGBColor(107, 114, 128)

        doc.add_page_break()

    # ── Front Matter: Copyright Page ──────────────────────────────────────────
    if ast.front_matter.copyright.enabled:
        c = ast.front_matter.copyright
        year = c.year or 2026
        holder = c.holder or ast.metadata.author

        # Vertical space push to bottom
        for _ in range(6):
            doc.add_paragraph()

        p_copy = doc.add_paragraph()
        p_copy.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_copy.paragraph_format.line_spacing = 1.15
        p_copy.paragraph_format.space_after = Pt(4)

        _add_formatted_text(p_copy, f"Copyright © {year} by {holder}\nAll rights reserved.")

        if c.statement:
            p_stmt = doc.add_paragraph()
            p_stmt.paragraph_format.space_after = Pt(4)
            _add_formatted_text(p_stmt, c.statement)

        if c.edition:
            p_ed = doc.add_paragraph()
            p_ed.paragraph_format.space_after = Pt(4)
            _add_formatted_text(p_ed, f"Edition: {c.edition}")

        if ast.metadata.isbn:
            p_isbn = doc.add_paragraph()
            p_isbn.paragraph_format.space_after = Pt(4)
            _add_formatted_text(p_isbn, f"ISBN: {ast.metadata.isbn}")

        if c.disclaimer:
            p_disc = doc.add_paragraph()
            p_disc.paragraph_format.space_after = Pt(4)
            _add_formatted_text(p_disc, c.disclaimer)

        doc.add_page_break()

    # ── Front Matter: Table of Contents ───────────────────────────────────────
    if ast.front_matter.table_of_contents.enabled:
        p_toc_title = doc.add_paragraph()
        p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_toc_title.paragraph_format.space_before = Pt(24)
        p_toc_title.paragraph_format.space_after = Pt(18)
        run_toc_title = p_toc_title.add_run(ast.front_matter.table_of_contents.title or "Table of Contents")
        run_toc_title.font.name = font_name
        run_toc_title.font.size = Pt(18)
        run_toc_title.bold = True

        for ch in ast.chapters:
            p_ch = doc.add_paragraph()
            p_ch.paragraph_format.space_after = Pt(6)
            p_ch.paragraph_format.line_spacing = 1.2
            
            run_ch_num = p_ch.add_run(f"Chapter {ch.chapter_number}:  ")
            run_ch_num.bold = True
            run_ch_num.font.size = Pt(10.5)

            run_ch_title = p_ch.add_run(ch.title)
            run_ch_title.font.size = Pt(10.5)

        doc.add_page_break()

    # ── Front Matter: Dedication ──────────────────────────────────────────────
    if ast.front_matter.dedication.enabled and ast.front_matter.dedication.text:
        for _ in range(4):
            doc.add_paragraph()

        p_ded = doc.add_paragraph()
        p_ded.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ded.paragraph_format.space_after = Pt(12)
        run_ded = p_ded.add_run(ast.front_matter.dedication.text)
        run_ded.font.name = font_name
        run_ded.font.size = Pt(12)
        run_ded.italic = True

        doc.add_page_break()

    # ── Chapters ──────────────────────────────────────────────────────────────
    for idx, chapter in enumerate(ast.chapters):
        if idx > 0 or ast.front_matter.title_page.enabled or ast.front_matter.copyright.enabled or ast.front_matter.table_of_contents.enabled:
            # New page for each chapter
            doc.add_page_break()

        # Chapter Number & Sinkage
        p_sink = doc.add_paragraph()
        p_sink.paragraph_format.space_before = Pt(36)
        p_sink.paragraph_format.space_after = Pt(6)
        run_ch_label = p_sink.add_run(f"CHAPTER {chapter.chapter_number}")
        run_ch_label.font.name = font_name
        run_ch_label.font.size = Pt(11)
        run_ch_label.font.color.rgb = RGBColor(107, 114, 128)
        run_ch_label.bold = True

        # Chapter Title
        p_ch_title = doc.add_paragraph()
        p_ch_title.paragraph_format.space_after = Pt(14)
        run_ch_main = p_ch_title.add_run(chapter.title)
        run_ch_main.font.name = font_name
        run_ch_main.font.size = Pt(22)
        run_ch_main.bold = True
        run_ch_main.font.color.rgb = RGBColor(17, 24, 39)

        # Chapter Subtitle
        if chapter.subtitle:
            p_ch_sub = doc.add_paragraph()
            p_ch_sub.paragraph_format.space_after = Pt(16)
            run_ch_sub = p_ch_sub.add_run(chapter.subtitle)
            run_ch_sub.font.name = font_name
            run_ch_sub.font.size = Pt(13)
            run_ch_sub.italic = True
            run_ch_sub.font.color.rgb = RGBColor(75, 85, 99)

        # Epigraph
        if chapter.epigraph:
            p_epi = doc.add_paragraph()
            p_epi.paragraph_format.left_indent = Inches(1.5)
            p_epi.paragraph_format.right_indent = Inches(0.25)
            p_epi.paragraph_format.space_before = Pt(8)
            p_epi.paragraph_format.space_after = Pt(20)
            
            run_epi = p_epi.add_run(f'"{chapter.epigraph.text}"')
            run_epi.font.name = font_name
            run_epi.font.size = Pt(10)
            run_epi.italic = True
            run_epi.font.color.rgb = RGBColor(75, 85, 99)

            if chapter.epigraph.attribution:
                p_epi_attr = doc.add_paragraph()
                p_epi_attr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p_epi_attr.paragraph_format.right_indent = Inches(0.25)
                p_epi_attr.paragraph_format.space_after = Pt(24)
                run_attr = p_epi_attr.add_run(f"— {chapter.epigraph.attribution}")
                run_attr.font.name = font_name
                run_attr.font.size = Pt(9.5)
                run_attr.font.color.rgb = RGBColor(107, 114, 128)

        # Content Blocks
        first_paragraph = True
        for block in chapter.content:
            if isinstance(block, ParagraphBlock):
                p = doc.add_paragraph()
                p.alignment = ALIGN_MAP.get(block.align, WD_ALIGN_PARAGRAPH.JUSTIFY)
                p.paragraph_format.line_spacing = cs.line_height or 1.35
                p.paragraph_format.space_after = Pt(6)

                # Indentation rule: first paragraph of chapter unindented, subsequent paragraphs indented if block.indent is True
                if block.indent and not first_paragraph:
                    p.paragraph_format.first_line_indent = Inches(0.25)
                else:
                    p.paragraph_format.first_line_indent = Inches(0.0)

                _add_formatted_text(p, block.text)
                first_paragraph = False

            elif isinstance(block, Heading2Block):
                p_h2 = doc.add_paragraph()
                p_h2.paragraph_format.space_before = Pt(16)
                p_h2.paragraph_format.space_after = Pt(6)
                p_h2.paragraph_format.keep_with_next = True
                run_h2 = p_h2.add_run(block.text)
                run_h2.font.name = font_name
                run_h2.font.size = Pt(14)
                run_h2.bold = True
                run_h2.font.color.rgb = RGBColor(31, 41, 55)
                first_paragraph = True

            elif isinstance(block, Heading3Block):
                p_h3 = doc.add_paragraph()
                p_h3.paragraph_format.space_before = Pt(12)
                p_h3.paragraph_format.space_after = Pt(4)
                p_h3.paragraph_format.keep_with_next = True
                run_h3 = p_h3.add_run(block.text)
                run_h3.font.name = font_name
                run_h3.font.size = Pt(12)
                run_h3.bold = True
                run_h3.font.color.rgb = RGBColor(55, 65, 81)
                first_paragraph = True

            elif isinstance(block, CalloutBlock):
                variant = block.variant or "info"
                fill_hex, border_hex, text_hex = CALLOUT_COLORS.get(variant, CALLOUT_COLORS["info"])

                # Create 1x1 table for styled callout box
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False

                cell = table.cell(0, 0)
                _set_cell_background(cell, fill_hex)
                _set_cell_borders(cell, left=border_hex, sz="24") # 3pt left border
                _set_cell_margins(cell, top=160, bottom=160, left=240, right=200)

                cell_para = cell.paragraphs[0]
                cell_para.paragraph_format.space_after = Pt(4)
                cell_para.paragraph_format.line_spacing = 1.25

                # Callout title with variant label
                variant_tag = f"[{variant.upper()}] "
                if block.title:
                    variant_tag += block.title
                
                run_vtag = cell_para.add_run(variant_tag + "\n")
                run_vtag.bold = True
                run_vtag.font.size = Pt(10)
                run_vtag.font.color.rgb = RGBColor(
                    int(text_hex[0:2], 16),
                    int(text_hex[2:4], 16),
                    int(text_hex[4:6], 16)
                )

                _add_formatted_text(cell_para, block.text)

                # Add small spacing after callout table
                p_spacer = doc.add_paragraph()
                p_spacer.paragraph_format.space_before = Pt(4)
                p_spacer.paragraph_format.space_after = Pt(4)
                first_paragraph = True

            elif isinstance(block, PullquoteBlock):
                p_pq = doc.add_paragraph()
                p_pq.alignment = ALIGN_MAP.get(block.align, WD_ALIGN_PARAGRAPH.CENTER)
                p_pq.paragraph_format.space_before = Pt(14)
                p_pq.paragraph_format.space_after = Pt(4)
                p_pq.paragraph_format.left_indent = Inches(0.5)
                p_pq.paragraph_format.right_indent = Inches(0.5)

                run_quote = p_pq.add_run(f'"{block.text}"')
                run_quote.font.name = font_name
                run_quote.font.size = Pt(13)
                run_quote.italic = True
                run_quote.font.color.rgb = RGBColor(75, 85, 99)

                if block.attribution:
                    p_attr = doc.add_paragraph()
                    p_attr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p_attr.paragraph_format.right_indent = Inches(0.5)
                    p_attr.paragraph_format.space_after = Pt(14)
                    run_pq_attr = p_attr.add_run(f"— {block.attribution}")
                    run_pq_attr.font.name = font_name
                    run_pq_attr.font.size = Pt(9.5)
                    run_pq_attr.font.color.rgb = RGBColor(107, 114, 128)

                first_paragraph = True

            elif isinstance(block, TableBlock):
                num_cols = len(block.headers)
                num_rows = len(block.rows) + 1
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = True

                # Header Row
                hdr_cells = table.rows[0].cells
                for col_idx, header_text in enumerate(block.headers):
                    cell = hdr_cells[col_idx]
                    _set_cell_background(cell, "374151") # Dark gray
                    _set_cell_margins(cell, top=140, bottom=140, left=160, right=160)
                    _set_cell_borders(cell, bottom="1F2937", sz="8")
                    p_cell = cell.paragraphs[0]
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_cell.paragraph_format.space_after = Pt(0)
                    run_h = p_cell.add_run(header_text)
                    run_h.bold = True
                    run_h.font.size = Pt(9.5)
                    run_h.font.color.rgb = RGBColor(255, 255, 255)

                # Data Rows
                for row_idx, row_data in enumerate(block.rows):
                    row_cells = table.rows[row_idx + 1].cells
                    is_striped = block.striped and (row_idx % 2 == 1)
                    bg_color = "F9FAFB" if is_striped else "FFFFFF"

                    for col_idx, cell_value in enumerate(row_data):
                        if col_idx < len(row_cells):
                            cell = row_cells[col_idx]
                            _set_cell_background(cell, bg_color)
                            _set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
                            _set_cell_borders(cell, bottom="E5E7EB", sz="4")
                            p_cell = cell.paragraphs[0]
                            align = block.column_alignments[col_idx] if block.column_alignments and col_idx < len(block.column_alignments) else "left"
                            p_cell.alignment = ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.LEFT)
                            p_cell.paragraph_format.space_after = Pt(0)
                            run_d = p_cell.add_run(str(cell_value))
                            run_d.font.size = Pt(9.5)
                            run_d.font.color.rgb = RGBColor(31, 41, 55)

                p_after_tbl = doc.add_paragraph()
                p_after_tbl.paragraph_format.space_after = Pt(8)
                first_paragraph = True

            elif isinstance(block, InteractiveFieldBlock):
                p_field = doc.add_paragraph()
                p_field.paragraph_format.space_before = Pt(8)
                p_field.paragraph_format.space_after = Pt(4)
                
                run_label = p_field.add_run(block.label)
                run_label.bold = True
                run_label.font.size = Pt(10)
                run_label.font.color.rgb = RGBColor(31, 41, 55)

                if block.field_type == "checkbox":
                    for opt in (block.options or ["Option 1", "Option 2"]):
                        p_opt = doc.add_paragraph()
                        p_opt.paragraph_format.left_indent = Inches(0.25)
                        p_opt.paragraph_format.space_after = Pt(2)
                        run_box = p_opt.add_run(f"☐  {opt}")
                        run_box.font.size = Pt(10)
                elif block.field_type in ("text", "multiline", "signature", "date"):
                    num_lines = block.lines if block.lines and block.lines > 0 else 2
                    for _ in range(num_lines):
                        p_line = doc.add_paragraph()
                        p_line.paragraph_format.space_after = Pt(4)
                        run_line = p_line.add_run("________________________________________________________")
                        run_line.font.color.rgb = RGBColor(156, 163, 175)

                first_paragraph = True

            elif isinstance(block, PageBreakBlock):
                doc.add_page_break()
                first_paragraph = True

            elif isinstance(block, HorizontalRuleBlock):
                p_hr = doc.add_paragraph()
                p_hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_hr.paragraph_format.space_before = Pt(10)
                p_hr.paragraph_format.space_after = Pt(10)
                run_hr = p_hr.add_run("*   *   *")
                run_hr.font.color.rgb = RGBColor(156, 163, 175)
                run_hr.font.size = Pt(10)
                first_paragraph = True

            elif isinstance(block, ImageBlock):
                p_img = doc.add_paragraph()
                p_img.alignment = ALIGN_MAP.get(block.align, WD_ALIGN_PARAGRAPH.CENTER)
                p_img.paragraph_format.space_before = Pt(8)
                p_img.paragraph_format.space_after = Pt(8)
                run_img = p_img.add_run(f"[ Image: {block.alt or block.caption or block.src} ]")
                run_img.italic = True
                run_img.font.color.rgb = RGBColor(107, 114, 128)
                first_paragraph = True

    return doc


class DocxCompiler(BaseCompiler):
    """Word (.docx) Compiler powered by python-docx."""

    @property
    def format_name(self) -> str:
        return "docx"

    @property
    def file_extension(self) -> str:
        return ".docx"

    @property
    def mime_type(self) -> str:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    async def compile(
        self,
        ast: DocumentAST,
        job_id: str,
        output_dir: Path,
    ) -> Tuple[Path, str]:
        output_dir.mkdir(parents=True, exist_ok=True)
        safe_title = sanitize_filename(ast.metadata.title)
        docx_filename = f"{job_id}_{safe_title}.docx"
        docx_path = output_dir / docx_filename

        # Build docx document in a thread to keep async event loop responsive
        def _build_and_save():
            doc = _build_docx_document(ast)
            doc.save(str(docx_path))

        try:
            await asyncio.to_thread(_build_and_save)
        except Exception as e:
            logger.error(f"DOCX compilation failed: {e}")
            raise RuntimeError(f"DOCX compiler failed: {e}")

        download_url = self.get_download_url(docx_filename)
        return docx_path, download_url
