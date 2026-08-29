"""
15-Page Limit Restriction Engine for BookCraft AI.
Implements 3-Stage Defense-in-Depth Restriction:
  - Stage 1: Ingestion Pre-Flight check & file slicing (PDF, DOCX, MD).
  - Stage 2: AST chapter & block token slicing with demo notice injection.
  - Stage 3: Output PDF physical capping with PyMuPDF (fitz) + Upsell Teaser page.
"""
from __future__ import annotations
import logging
import os
import shutil
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import fitz  # PyMuPDF
import docx

from app.core.config import settings
from app.models.document_ast import (
    CalloutBlock,
    Chapter,
    DocumentAST,
    ParagraphBlock,
)

logger = logging.getLogger("bookcraft.restriction_engine")

DEMO_MAX_PAGES = 15
DEMO_MAX_WORDS = 4500
WORDS_PER_PAGE_ESTIMATE = 300


@dataclass
class PreflightResult:
    """Result of Stage 1 Ingestion Pre-Flight check."""
    is_truncated: bool
    original_pages: Optional[int] = None
    original_words: Optional[int] = None
    sliced_pages: Optional[int] = None
    sliced_words: Optional[int] = None
    file_path: str = ""
    message: str = ""

    def to_dict(self) -> Dict[str, Any]:
        return {
            "is_truncated": self.is_truncated,
            "original_pages": self.original_pages,
            "original_words": self.original_words,
            "sliced_pages": self.sliced_pages,
            "sliced_words": self.sliced_words,
            "file_path": self.file_path,
            "message": self.message,
        }


@dataclass
class OutputCappingResult:
    """Result of Stage 3 Output PDF physical capping."""
    is_capped: bool
    original_page_count: int
    final_page_count: int
    pdf_path: str
    upsell_page_appended: bool = False
    message: str = ""

    def to_dict(self) -> Dict[str, Any]:
        return {
            "is_capped": self.is_capped,
            "original_page_count": self.original_page_count,
            "final_page_count": self.final_page_count,
            "pdf_path": self.pdf_path,
            "upsell_page_appended": self.upsell_page_appended,
            "message": self.message,
        }


# ─────────────────────────────────────────────────────────────────────────────
# STAGE 1: Ingestion Pre-Flight Check & File Slicing
# ─────────────────────────────────────────────────────────────────────────────

def preflight_check_and_slice(
    file_path: str,
    file_type: str,
    is_demo: bool = True,
    max_pages: int = DEMO_MAX_PAGES,
    max_words: int = DEMO_MAX_WORDS,
) -> PreflightResult:
    """
    Inspects uploaded file dimensions (pages for PDF, words for DOCX/MD).
    If file exceeds demo tier limits, generates a sliced copy and returns metadata.
    """
    path = Path(file_path)
    if not path.exists():
        return PreflightResult(is_truncated=False, file_path=file_path, message="File not found")

    if not is_demo:
        return PreflightResult(
            is_truncated=False,
            file_path=file_path,
            message="Pro tier — no demo restrictions applied.",
        )

    file_ext = file_type.lower().lstrip(".")

    try:
        if file_ext == "pdf":
            return _preflight_slice_pdf(path, max_pages)
        elif file_ext in ("docx", "doc"):
            return _preflight_slice_docx(path, max_words)
        elif file_ext in ("md", "markdown", "txt"):
            return _preflight_slice_markdown(path, max_words)
        else:
            return PreflightResult(is_truncated=False, file_path=file_path, message="Unrestricted format")
    except Exception as exc:
        logger.warning(f"Preflight slicing encountered non-fatal error: {exc}. Proceeding with original file.")
        return PreflightResult(is_truncated=False, file_path=file_path, message=f"Preflight error: {exc}")


def _preflight_slice_pdf(path: Path, max_pages: int) -> PreflightResult:
    """Count PDF pages using PyMuPDF and slice to max_pages if exceeded."""
    doc = fitz.open(str(path))
    total_pages = len(doc)

    # Estimate word count
    total_words = 0
    for page in doc:
        total_words += len(page.get_text().split())

    if total_pages <= max_pages:
        doc.close()
        return PreflightResult(
            is_truncated=False,
            original_pages=total_pages,
            original_words=total_words,
            sliced_pages=total_pages,
            sliced_words=total_words,
            file_path=str(path),
            message=f"PDF has {total_pages} pages (within {max_pages} page demo limit).",
        )

    # File exceeds limit — slice to max_pages
    sliced_path = path.parent / f"{path.stem}_demo_sliced.pdf"
    sliced_doc = fitz.open()
    sliced_doc.insert_pdf(doc, from_page=0, to_page=max_pages - 1)
    sliced_doc.save(str(sliced_path))

    sliced_words = sum(len(page.get_text().split()) for page in sliced_doc)

    doc.close()
    sliced_doc.close()

    logger.info(
        f"[Stage 1] Truncated PDF '{path.name}' from {total_pages} pages to {max_pages} pages -> '{sliced_path.name}'"
    )
    return PreflightResult(
        is_truncated=True,
        original_pages=total_pages,
        original_words=total_words,
        sliced_pages=max_pages,
        sliced_words=sliced_words,
        file_path=str(sliced_path),
        message=f"Demo tier limit applied: Truncated from {total_pages} to {max_pages} pages.",
    )


def _preflight_slice_docx(path: Path, max_words: int) -> PreflightResult:
    """Count words in DOCX and slice paragraphs if limit exceeded."""
    try:
        doc = docx.Document(str(path))
    except Exception as exc:
        logger.warning(f"Could not open docx with python-docx: {exc}")
        return PreflightResult(is_truncated=False, file_path=str(path))

    # Calculate total word count
    total_words = 0
    for p in doc.paragraphs:
        total_words += len(p.text.split())
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                total_words += len(cell.text.split())

    est_pages = max(1, round(total_words / WORDS_PER_PAGE_ESTIMATE))

    if total_words <= max_words:
        return PreflightResult(
            is_truncated=False,
            original_pages=est_pages,
            original_words=total_words,
            sliced_pages=est_pages,
            sliced_words=total_words,
            file_path=str(path),
            message=f"DOCX has ~{total_words} words (~{est_pages} pages, within demo limit).",
        )

    # Slice docx
    sliced_doc = docx.Document()
    cumulative_words = 0
    for p in doc.paragraphs:
        words_in_p = len(p.text.split())
        if cumulative_words + words_in_p > max_words and cumulative_words > 0:
            break
        new_p = sliced_doc.add_paragraph(p.text, style=p.style.name if p.style else None)
        cumulative_words += words_in_p

    # Add notice paragraph
    notice = sliced_doc.add_paragraph()
    run = notice.add_run(
        "\n[Demo Preview Notice: Manuscript truncated to first 15 pages (4,500 words). "
        "Upgrade to BookCraft Pro to format and compile your entire book.]"
    )
    run.italic = True

    sliced_path = path.parent / f"{path.stem}_demo_sliced.docx"
    sliced_doc.save(str(sliced_path))

    logger.info(
        f"[Stage 1] Truncated DOCX '{path.name}' from {total_words} words to {cumulative_words} words -> '{sliced_path.name}'"
    )
    return PreflightResult(
        is_truncated=True,
        original_pages=est_pages,
        original_words=total_words,
        sliced_pages=DEMO_MAX_PAGES,
        sliced_words=cumulative_words,
        file_path=str(sliced_path),
        message=f"Demo tier limit applied: Truncated from ~{total_words} words to {cumulative_words} words.",
    )


def _preflight_slice_markdown(path: Path, max_words: int) -> PreflightResult:
    """Count words in Markdown and slice text to max_words."""
    text = path.read_text(encoding="utf-8", errors="ignore")
    words = text.split()
    total_words = len(words)
    est_pages = max(1, round(total_words / WORDS_PER_PAGE_ESTIMATE))

    if total_words <= max_words:
        return PreflightResult(
            is_truncated=False,
            original_pages=est_pages,
            original_words=total_words,
            sliced_pages=est_pages,
            sliced_words=total_words,
            file_path=str(path),
            message=f"Markdown has ~{total_words} words (within demo limit).",
        )

    # Slices at nearest paragraph break within limit
    paragraphs = text.split("\n\n")
    sliced_paragraphs = []
    accumulated_words = 0

    for para in paragraphs:
        p_words = len(para.split())
        if accumulated_words + p_words > max_words and accumulated_words > 0:
            break
        sliced_paragraphs.append(para)
        accumulated_words += p_words

    sliced_paragraphs.append(
        "\n> **Demo Preview**: Manuscript truncated to 15 pages (~4,500 words). "
        "Upgrade to BookCraft Pro to format your complete manuscript.\n"
    )

    sliced_content = "\n\n".join(sliced_paragraphs)
    sliced_path = path.parent / f"{path.stem}_demo_sliced.md"
    sliced_path.write_text(sliced_content, encoding="utf-8")

    logger.info(
        f"[Stage 1] Truncated Markdown '{path.name}' from {total_words} words to {accumulated_words} words."
    )
    return PreflightResult(
        is_truncated=True,
        original_pages=est_pages,
        original_words=total_words,
        sliced_pages=DEMO_MAX_PAGES,
        sliced_words=accumulated_words,
        file_path=str(sliced_path),
        message=f"Demo tier limit applied: Truncated from ~{total_words} to {accumulated_words} words.",
    )


# ─────────────────────────────────────────────────────────────────────────────
# STAGE 2: AST Chapter & Content Slicing for Demo Tier
# ─────────────────────────────────────────────────────────────────────────────

def slice_ast_for_demo(
    ast: DocumentAST,
    is_demo: bool = True,
    max_words: int = DEMO_MAX_WORDS,
) -> DocumentAST:
    """
    Slices DocumentAST chapters and blocks to stay strictly within demo constraints.
    Appends a formatted demo limitation callout block at the end.
    """
    if not is_demo:
        return ast

    total_words = sum(c.word_count or 0 for c in ast.chapters)
    if total_words <= max_words and len(ast.chapters) <= 3:
        # If document already fits well within 15 pages, return as is
        return ast

    sliced_chapters: List[Chapter] = []
    cumulative_words = 0

    for chapter in ast.chapters:
        ch_words = chapter.word_count or sum(
            len(b.text.split()) for b in chapter.content if hasattr(b, "text") and b.text
        )

        if cumulative_words + ch_words <= max_words:
            sliced_chapters.append(chapter)
            cumulative_words += ch_words
        else:
            # Need to slice this chapter's blocks
            sliced_blocks = []
            block_words = 0
            for block in chapter.content:
                b_text = getattr(block, "text", "") or ""
                b_count = len(b_text.split())
                if cumulative_words + block_words + b_count > max_words and (sliced_blocks or sliced_chapters):
                    break
                sliced_blocks.append(block)
                block_words += b_count

            # Add demo teaser callout block
            demo_notice = CalloutBlock(
                callout_type="tip",
                title="Demo Preview — 15-Page Limit Reached",
                content=(
                    "You are viewing a free 15-page sample formatted by BookCraft AI. "
                    "Upgrade to BookCraft Pro to format, compile, and download your entire book "
                    "with all chapters, full table of contents, and editable DOCX/Markdown files."
                ),
            )
            sliced_blocks.append(demo_notice)

            chapter.content = sliced_blocks
            chapter.word_count = block_words
            sliced_chapters.append(chapter)
            cumulative_words += block_words
            break  # Do not include subsequent chapters

    if not sliced_chapters and ast.chapters:
        sliced_chapters = [ast.chapters[0]]

    ast.chapters = sliced_chapters
    logger.info(
        f"[Stage 2] Sliced AST for demo: {len(sliced_chapters)} chapters retained, ~{cumulative_words} words."
    )
    return ast


# ─────────────────────────────────────────────────────────────────────────────
# STAGE 3: PyMuPDF Output PDF Capping & Upsell Page Insertion
# ─────────────────────────────────────────────────────────────────────────────

def enforce_output_pdf_limit(
    pdf_path: str,
    is_demo: bool = True,
    book_title: str = "Your Book",
    author_name: str = "Author",
    max_pages: int = DEMO_MAX_PAGES,
) -> OutputCappingResult:
    """
    Opens compiled PDF with PyMuPDF (fitz). If page count > max_pages (15) and is_demo:
      1. Strictly truncates PDF to exactly 15 pages.
      2. Generates and appends a stylized Page 16 'Demo Upsell & Upgrade' teaser card.
      3. Saves modified PDF in place.
    """
    path = Path(pdf_path)
    if not path.exists():
        return OutputCappingResult(
            is_capped=False,
            original_page_count=0,
            final_page_count=0,
            pdf_path=pdf_path,
            message="PDF file not found.",
        )

    if not is_demo:
        try:
            doc = fitz.open(str(path))
            p_count = len(doc)
            doc.close()
        except Exception:
            p_count = 0
        return OutputCappingResult(
            is_capped=False,
            original_page_count=p_count,
            final_page_count=p_count,
            pdf_path=pdf_path,
            message="Pro tier — no output truncation applied.",
        )

    doc = fitz.open(str(path))
    orig_pages = len(doc)

    if orig_pages <= max_pages:
        doc.close()
        return OutputCappingResult(
            is_capped=False,
            original_page_count=orig_pages,
            final_page_count=orig_pages,
            pdf_path=pdf_path,
            upsell_page_appended=False,
            message=f"PDF fits within demo limit ({orig_pages} pages <= {max_pages}).",
        )

    # ── Strict Capping to max_pages + Upsell Page ────────────────────────────
    capped_doc = fitz.open()
    # Insert first 15 pages (index 0 to max_pages - 1)
    capped_doc.insert_pdf(doc, from_page=0, to_page=max_pages - 1)

    # Get dimensions from first page or default to 6x9 (432 x 648 pt)
    ref_rect = doc[0].rect if orig_pages > 0 else fitz.Rect(0, 0, 432, 648)
    width, height = ref_rect.width, ref_rect.height

    # Append stylized Upsell Page (Page 16)
    upsell_page = capped_doc.new_page(width=width, height=height)
    _render_upsell_page(upsell_page, width, height, book_title, author_name, orig_pages)

    # Save to temp file and atomically replace
    temp_path = path.parent / f"{path.stem}_capped_tmp.pdf"
    capped_doc.save(str(temp_path), garbage=4, deflate=True)

    doc.close()
    capped_doc.close()

    shutil.move(str(temp_path), str(path))
    logger.info(
        f"[Stage 3] Enforced demo PDF capping on '{path.name}': "
        f"Truncated from {orig_pages} pages -> {max_pages} pages + Upsell page appended."
    )

    return OutputCappingResult(
        is_capped=True,
        original_page_count=orig_pages,
        final_page_count=max_pages + 1,
        pdf_path=str(path),
        upsell_page_appended=True,
        message=f"Truncated from {orig_pages} to {max_pages} pages with Pro upgrade invitation.",
    )


def _render_upsell_page(
    page: fitz.Page,
    width: float,
    height: float,
    book_title: str,
    author_name: str,
    original_pages: int,
) -> None:
    """Draws a professional, clean BookCraft Pro upsell teaser page in PyMuPDF."""
    # Background soft color & border
    bg_rect = fitz.Rect(20, 20, width - 20, height - 20)
    page.draw_rect(bg_rect, color=(0.85, 0.88, 0.95), fill=(0.98, 0.98, 1.0), width=1.5)

    # Header Badge Box
    badge_rect = fitz.Rect(width / 2 - 90, 45, width / 2 + 90, 75)
    page.draw_rect(badge_rect, color=(0.2, 0.4, 0.8), fill=(0.2, 0.4, 0.8))
    page.insert_textbox(
        badge_rect,
        "BOOKCRAFT AI — PRO PREVIEW",
        fontsize=9,
        color=(1, 1, 1),
        align=fitz.TEXT_ALIGN_CENTER,
    )

    # Title & Subtitle Text
    title_rect = fitz.Rect(35, 90, width - 35, 140)
    page.insert_textbox(
        title_rect,
        "End of Free 15-Page Preview",
        fontsize=16,
        color=(0.1, 0.15, 0.3),
        align=fitz.TEXT_ALIGN_CENTER,
    )

    meta_rect = fitz.Rect(35, 145, width - 35, 185)
    page.insert_textbox(
        meta_rect,
        f"Manuscript: \"{book_title[:45]}\"\nOriginal Length: ~{original_pages} pages",
        fontsize=10,
        color=(0.35, 0.4, 0.5),
        align=fitz.TEXT_ALIGN_CENTER,
    )

    # Decorative Line
    page.draw_line(
        fitz.Point(45, 195),
        fitz.Point(width - 45, 195),
        color=(0.8, 0.85, 0.9),
        width=1,
    )

    # Feature List Box
    features_rect = fitz.Rect(45, 210, width - 45, 380)
    features_text = (
        "Unlock the Full Commercial BookCraft Pro Suite:\n\n"
        "  [✓] Complete Manuscript Formatting (Unlimited Pages)\n"
        "  [✓] High-Resolution 300 DPI Print-Ready PDF (KDP / IngramSpark)\n"
        "  [✓] Editable DOCX, Markdown (.md), and EPUB3 Digital Editions\n"
        "  [✓] Custom Typography, Drop Caps, Callouts & Chapter Layouts\n"
        "  [✓] Zero Watermarks & 100% Commercial Publishing Rights\n"
    )
    page.insert_textbox(
        features_rect,
        features_text,
        fontsize=9.5,
        color=(0.15, 0.2, 0.3),
        align=fitz.TEXT_ALIGN_LEFT,
    )

    # Call to Action Card
    cta_rect = fitz.Rect(40, height - 140, width - 40, height - 40)
    page.draw_rect(cta_rect, color=(0.15, 0.35, 0.75), fill=(0.93, 0.96, 1.0), width=1.2)

    cta_text = (
        "Ready to Publish Your Full Book?\n\n"
        "Upgrade now in the BookCraft AI Editor by clicking 'Unlock Full Book'\n"
        "or visit: https://bookcraft.ai/checkout"
    )
    page.insert_textbox(
        cta_rect,
        cta_text,
        fontsize=9.5,
        color=(0.1, 0.2, 0.5),
        align=fitz.TEXT_ALIGN_CENTER,
    )
