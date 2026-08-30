"""
Unit tests for modular multi-format compilers (PDF, DOCX, Markdown, EPUB).
"""
import io
import os
from pathlib import Path
import tempfile
import zipfile
import pytest

from app.models.document_ast import (
    BookMetadata,
    CalloutBlock,
    Chapter,
    CompilationSettings,
    CopyrightPage,
    DedicationPage,
    DocumentAST,
    Epigraph,
    FrontMatter,
    Genre,
    Heading2Block,
    Heading3Block,
    HorizontalRuleBlock,
    ImageBlock,
    InteractiveFieldBlock,
    PageBreakBlock,
    ParagraphBlock,
    PullquoteBlock,
    TableBlock,
    TableOfContents,
    TitlePage,
    TrimSize,
)
from app.services.compilers import (
    DocxCompiler,
    EpubCompiler,
    MdCompiler,
    PdfCompiler,
    CompilerOrchestrator,
    compile_all_formats,
)
from tests.conftest import (
    is_valid_docx_bytes,
    is_valid_epub_bytes,
    is_valid_md_string,
    is_valid_pdf_bytes,
)


def get_rich_ast() -> DocumentAST:
    """DocumentAST containing every block type and front matter feature."""
    return DocumentAST(
        metadata=BookMetadata(
            title="Mastering BookCraft AI",
            subtitle="The Complete Architecture & Implementation Manual",
            author="Dr. Alex Rivera",
            co_authors=["Elena Rostova", "Kai Tanaka"],
            genre=Genre.technical,
            trim_size=TrimSize.medium,
            isbn="978-1-987654-32-1",
            publisher="Algorithmic Publishing House",
            published_year=2026,
            edition="Second Edition",
            language="en",
            keywords=["compilers", "fastapi", "nextjs", "typst", "epub", "docx"],
        ),
        front_matter=FrontMatter(
            title_page=TitlePage(
                enabled=True,
                display_title="Mastering BookCraft AI",
                display_subtitle="The Complete Architecture & Implementation Manual",
                display_author="Dr. Alex Rivera",
                display_publisher="Algorithmic Publishing House",
            ),
            copyright=CopyrightPage(
                enabled=True,
                year=2026,
                holder="Dr. Alex Rivera",
                statement="All rights reserved under International Copyright Conventions.",
                edition="Second Edition",
                disclaimer="Software code examples provided AS-IS.",
            ),
            table_of_contents=TableOfContents(
                enabled=True,
                title="Table of Contents",
                include_subheadings=True,
                max_depth=2,
            ),
            dedication=DedicationPage(
                enabled=True,
                text="Dedicated to all open source contributors building the future of documents.",
            ),
        ),
        chapters=[
            Chapter(
                chapter_number=1,
                title="Architecture & Multi-Format Pipeline",
                subtitle="From AST to Print, Web, Word, and E-Reader",
                epigraph=Epigraph(
                    text="The ultimate goal of typesetting is transparency to thought.",
                    attribution="Classic Typographic Principle",
                ),
                content=[
                    ParagraphBlock(
                        type="paragraph",
                        text="This chapter introduces the **modular compiler architecture** of BookCraft AI. We combine *high-performance formatting* with ***pristine semantic clarity*** and `clean code` interfaces.",
                        indent=True,
                        align="justify",
                    ),
                    Heading2Block(
                        type="heading2",
                        text="Compiler Abstraction Layer",
                    ),
                    Heading3Block(
                        type="heading3",
                        text="Unified Interface Protocol",
                    ),
                    CalloutBlock(
                        type="callout",
                        variant="tip",
                        title="Pro Developer Tip",
                        text="Always use BaseCompiler protocol to plug in new formats effortlessly.",
                    ),
                    CalloutBlock(
                        type="callout",
                        variant="warning",
                        title="Cautionary Note",
                        text="Ensure all file paths are validated against directory traversal attacks.",
                    ),
                    PullquoteBlock(
                        type="pullquote",
                        text="A truly great publishing engine treats every output format as a first-class citizen.",
                        attribution="Chief Architect",
                        align="center",
                    ),
                    TableBlock(
                        type="table",
                        caption="Format Capability Comparison",
                        headers=["Format", "Extension", "Engine", "Reflowable"],
                        rows=[
                            ["PDF", ".pdf", "Typst", "No (Fixed Layout)"],
                            ["Word", ".docx", "python-docx", "Yes (Page Layout)"],
                            ["Markdown", ".md", "Native Python", "Yes (Plain Text)"],
                            ["EPUB", ".epub", "EPUB3 Zip Packager", "Yes (Reflowable)"],
                        ],
                        column_alignments=["left", "center", "left", "center"],
                        striped=True,
                    ),
                    InteractiveFieldBlock(
                        type="interactive-field",
                        field_type="checkbox",
                        label="Select Target Formats:",
                        options=["PDF Export", "DOCX Export", "Markdown Export", "EPUB Export"],
                    ),
                    InteractiveFieldBlock(
                        type="interactive-field",
                        field_type="text",
                        label="Author Notes & Revisions",
                        lines=3,
                    ),
                    HorizontalRuleBlock(
                        type="horizontal-rule",
                        style="dots",
                    ),
                    PageBreakBlock(
                        type="page-break",
                    ),
                    ImageBlock(
                        type="image",
                        src="architecture_diagram.png",
                        alt="BookCraft Architecture Diagram",
                        caption="Figure 1.1: Multi-format compilation flow",
                    ),
                    ParagraphBlock(
                        type="paragraph",
                        text="Concluding the architectural overview with final verification notes.",
                        indent=True,
                        align="justify",
                    ),
                ],
                word_count=650,
            ),
            Chapter(
                chapter_number=2,
                title="Production Deployment & Verification",
                subtitle="Ensuring Zero-Defect Artifact Generation",
                content=[
                    ParagraphBlock(
                        type="paragraph",
                        text="Continuous integration suites verify every generated binary format against standard specifications.",
                        indent=True,
                        align="justify",
                    ),
                ],
                word_count=300,
            ),
        ],
        compilation_settings=CompilationSettings(
            trim_size=TrimSize.medium,
            font_size=11,
            line_height=1.5,
        ),
    )


@pytest.fixture
def rich_ast() -> DocumentAST:
    return get_rich_ast()


@pytest.mark.asyncio
async def test_pdf_compiler_generates_valid_pdf(sample_ast, tmp_path):
    """Verify PdfCompiler generates non-empty valid PDF with %PDF header."""
    compiler = PdfCompiler()
    assert compiler.format_name == "pdf"
    assert compiler.file_extension == ".pdf"
    assert compiler.mime_type == "application/pdf"

    job_id = "test_job_pdf_001"
    output_path, download_url = await compiler.compile(sample_ast, job_id, tmp_path)

    assert output_path.exists()
    assert output_path.stat().st_size > 0
    assert download_url == f"/api/download/{output_path.name}"

    pdf_bytes = output_path.read_bytes()
    assert is_valid_pdf_bytes(pdf_bytes)


@pytest.mark.asyncio
async def test_docx_compiler_generates_valid_docx(rich_ast, tmp_path):
    """Verify DocxCompiler generates valid Word docx with all block types preserved."""
    compiler = DocxCompiler()
    assert compiler.format_name == "docx"
    assert compiler.file_extension == ".docx"
    assert "wordprocessingml" in compiler.mime_type

    job_id = "test_job_docx_001"
    output_path, download_url = await compiler.compile(rich_ast, job_id, tmp_path)

    assert output_path.exists()
    assert output_path.stat().st_size > 0
    assert download_url == f"/api/download/{output_path.name}"

    docx_bytes = output_path.read_bytes()
    assert is_valid_docx_bytes(docx_bytes)

    # Verify internal docx structure
    from docx import Document
    doc = Document(str(output_path))
    doc_text = " ".join([p.text for p in doc.paragraphs])
    assert "Mastering BookCraft AI" in doc_text
    assert "Dr. Alex Rivera" in doc_text
    assert "Architecture & Multi-Format Pipeline" in doc_text

    # Verify table existence in docx
    assert len(doc.tables) >= 2  # At least the callout table and the comparison table


@pytest.mark.asyncio
async def test_md_compiler_generates_clean_markdown(rich_ast, tmp_path):
    """Verify MdCompiler generates clean markdown with YAML frontmatter, tables, and GFM callouts."""
    compiler = MdCompiler()
    assert compiler.format_name == "md"
    assert compiler.file_extension == ".md"
    assert "text/markdown" in compiler.mime_type

    job_id = "test_job_md_001"
    output_path, download_url = await compiler.compile(rich_ast, job_id, tmp_path)

    assert output_path.exists()
    assert output_path.stat().st_size > 0
    assert download_url == f"/api/download/{output_path.name}"

    md_content = output_path.read_text(encoding="utf-8")
    assert is_valid_md_string(md_content)

    # Verify YAML frontmatter
    assert md_content.startswith("---")
    assert 'title: "Mastering BookCraft AI"' in md_content
    assert 'author: "Dr. Alex Rivera"' in md_content
    assert 'genre: "technical"' in md_content

    # Verify GFM elements
    assert "# Chapter 1: Architecture & Multi-Format Pipeline" in md_content
    assert "> [!TIP]" in md_content
    assert "> [!WARNING]" in md_content
    assert "| Format | Extension | Engine | Reflowable |" in md_content
    assert "- [ ] PDF Export" in md_content
    assert "<!-- pagebreak -->" in md_content


@pytest.mark.asyncio
async def test_epub_compiler_generates_valid_epub3(rich_ast, tmp_path):
    """Verify EpubCompiler generates standard EPUB3 package with mimetype, nav, opf, and styles."""
    compiler = EpubCompiler()
    assert compiler.format_name == "epub"
    assert compiler.file_extension == ".epub"
    assert compiler.mime_type == "application/epub+zip"

    job_id = "test_job_epub_001"
    output_path, download_url = await compiler.compile(rich_ast, job_id, tmp_path)

    assert output_path.exists()
    assert output_path.stat().st_size > 0
    assert download_url == f"/api/download/{output_path.name}"

    epub_bytes = output_path.read_bytes()
    assert is_valid_epub_bytes(epub_bytes)

    # Verify internal zip archive structure
    with zipfile.ZipFile(output_path, "r") as zf:
        namelist = zf.namelist()
        # mimetype MUST be first file in archive
        assert namelist[0] == "mimetype"
        assert zf.read("mimetype") == b"application/epub+zip"

        # Check required EPUB3 components
        assert "META-INF/container.xml" in namelist
        assert "OEBPS/content.opf" in namelist
        assert "OEBPS/nav.xhtml" in namelist
        assert "OEBPS/toc.ncx" in namelist
        assert "OEBPS/style.css" in namelist
        assert "OEBPS/chapter_1.xhtml" in namelist
        assert "OEBPS/chapter_2.xhtml" in namelist

        # Inspect chapter 1 content
        ch1_content = zf.read("OEBPS/chapter_1.xhtml").decode("utf-8")
        assert "Architecture &amp; Multi-Format Pipeline" in ch1_content or "Architecture & Multi-Format Pipeline" in ch1_content
        assert "callout-tip" in ch1_content
        assert "<table>" in ch1_content


@pytest.mark.asyncio
async def test_compiler_orchestrator_compiles_all_formats(rich_ast, tmp_path):
    """Verify CompilerOrchestrator compiles all 4 formats concurrently and returns complete metadata."""
    orchestrator = CompilerOrchestrator()
    job_id = "test_job_orch_001"

    results = await orchestrator.compile_all(
        ast=rich_ast,
        job_id=job_id,
        output_dir=tmp_path,
    )

    assert "pdf" in results
    assert "docx" in results
    assert "md" in results
    assert "epub" in results

    for fmt in ["pdf", "docx", "md", "epub"]:
        data = results[fmt]
        path = Path(data["path"])
        assert path.exists(), f"File for {fmt} was not created"
        assert path.stat().st_size > 0, f"File for {fmt} is empty"
        assert data["url"] == f"/api/download/{path.name}"
        assert data["size_bytes"] == path.stat().st_size


@pytest.mark.asyncio
async def test_unicode_and_special_characters_across_all_formats(unicode_ast, tmp_path):
    """Verify all 4 compilers handle unicode, emojis, accents, and non-latin characters cleanly."""
    job_id = "test_job_unicode_001"
    results = await compile_all_formats(
        ast=unicode_ast,
        job_id=job_id,
        output_dir=tmp_path,
    )

    assert len(results) == 4
    for fmt, data in results.items():
        p = Path(data["path"])
        assert p.exists()
        assert p.stat().st_size > 0
