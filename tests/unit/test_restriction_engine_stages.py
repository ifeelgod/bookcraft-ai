"""
Unit Tests: 3-Stage Restriction Engine Implementation
Directly tests:
  - Stage 1: Ingestion preflight slicing (PDF, DOCX, Markdown)
  - Stage 2: AST chapter & content token slicing
  - Stage 3: Output PDF physical capping with PyMuPDF & Page 16 upsell card insertion
"""
from __future__ import annotations
import copy
import tempfile
from pathlib import Path
import pytest
import fitz
import docx

from app.models.document_ast import (
    BookMetadata,
    CalloutBlock,
    Chapter,
    DocumentAST,
    Genre,
    ParagraphBlock,
    TrimSize,
)
from app.services.restriction_engine import (
    DEMO_MAX_PAGES,
    DEMO_MAX_WORDS,
    preflight_check_and_slice,
    slice_ast_for_demo,
    enforce_output_pdf_limit,
    PreflightResult,
    OutputCappingResult,
)


@pytest.fixture
def temp_dir():
    with tempfile.TemporaryDirectory() as tmpdir:
        yield Path(tmpdir)


def create_test_pdf(path: Path, num_pages: int = 20) -> Path:
    """Helper to create a synthetic multi-page PDF using PyMuPDF."""
    doc = fitz.open()
    for i in range(num_pages):
        page = doc.new_page(width=432, height=648)
        page.insert_text((50, 100), f"Chapter {i + 1} - Page Content", fontsize=14)
        for line in range(10):
            page.insert_text((50, 140 + line * 20), f"Line {line + 1}: The narrative unfolds across page {i + 1}.", fontsize=10)
    doc.save(str(path))
    doc.close()
    return path


def create_test_docx(path: Path, num_words: int = 6000) -> Path:
    """Helper to create a synthetic DOCX file with specified word count."""
    doc = docx.Document()
    words_per_para = 50
    num_paras = num_words // words_per_para
    for i in range(num_paras):
        doc.add_paragraph(f"Paragraph {i + 1}: " + " ".join(["word"] * (words_per_para - 2)) + ".")
    doc.save(str(path))
    return path


def create_test_md(path: Path, num_words: int = 6000) -> Path:
    """Helper to create a synthetic Markdown file with specified word count."""
    paragraphs = []
    words_per_para = 100
    num_paras = num_words // words_per_para
    for i in range(num_paras):
        paragraphs.append(f"## Section {i + 1}\n\n" + " ".join(["lore"] * (words_per_para - 2)) + ".")
    path.write_text("\n\n".join(paragraphs), encoding="utf-8")
    return path


# ── Stage 1 Tests ─────────────────────────────────────────────────────────────

def test_stage1_pdf_preflight_slice_exceeds_limit(temp_dir):
    """Verify that a 20-page PDF is sliced down to 15 pages in demo mode."""
    pdf_path = temp_dir / "manuscript_20p.pdf"
    create_test_pdf(pdf_path, num_pages=20)

    result: PreflightResult = preflight_check_and_slice(
        file_path=str(pdf_path),
        file_type="pdf",
        is_demo=True,
    )

    assert result.is_truncated is True
    assert result.original_pages == 20
    assert result.sliced_pages == 15
    assert Path(result.file_path).exists()
    assert "_demo_sliced" in result.file_path

    # Verify the sliced PDF actually has 15 pages
    doc = fitz.open(result.file_path)
    assert len(doc) == 15
    doc.close()


def test_stage1_pdf_preflight_short_doc_not_sliced(temp_dir):
    """Verify that a 10-page PDF is not sliced."""
    pdf_path = temp_dir / "short_10p.pdf"
    create_test_pdf(pdf_path, num_pages=10)

    result = preflight_check_and_slice(
        file_path=str(pdf_path),
        file_type="pdf",
        is_demo=True,
    )

    assert result.is_truncated is False
    assert result.original_pages == 10
    assert result.file_path == str(pdf_path)


def test_stage1_pdf_preflight_pro_tier_bypasses(temp_dir):
    """Verify that Pro tier preserves all 30 pages without slicing."""
    pdf_path = temp_dir / "epic_30p.pdf"
    create_test_pdf(pdf_path, num_pages=30)

    result = preflight_check_and_slice(
        file_path=str(pdf_path),
        file_type="pdf",
        is_demo=False,
    )

    assert result.is_truncated is False
    assert result.file_path == str(pdf_path)


def test_stage1_docx_preflight_slice(temp_dir):
    """Verify that a 6,000-word DOCX is sliced to <= 4,500 words in demo mode."""
    docx_path = temp_dir / "manuscript_6000w.docx"
    create_test_docx(docx_path, num_words=6000)

    result = preflight_check_and_slice(
        file_path=str(docx_path),
        file_type="docx",
        is_demo=True,
    )

    assert result.is_truncated is True
    assert result.original_words >= 5500
    assert result.sliced_words <= DEMO_MAX_WORDS
    assert Path(result.file_path).exists()

    # Open sliced docx and verify notice paragraph
    doc = docx.Document(result.file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Demo Preview Notice" in full_text


def test_stage1_markdown_preflight_slice(temp_dir):
    """Verify that a 7,000-word Markdown file is sliced to <= 4,500 words in demo mode."""
    md_path = temp_dir / "manuscript_7000w.md"
    create_test_md(md_path, num_words=7000)

    result = preflight_check_and_slice(
        file_path=str(md_path),
        file_type="md",
        is_demo=True,
    )

    assert result.is_truncated is True
    assert result.sliced_words <= DEMO_MAX_WORDS
    assert Path(result.file_path).exists()

    content = Path(result.file_path).read_text(encoding="utf-8")
    assert "Demo Preview" in content


# ── Stage 2 Tests ─────────────────────────────────────────────────────────────

def test_stage2_ast_slicing_and_callout_injection(twenty_five_page_ast):
    """Verify AST chapter/content token slicing and demo notice callout block injection."""
    ast_copy = copy.deepcopy(twenty_five_page_ast)
    # Give chapters substantial word count
    for ch in ast_copy.chapters:
        ch.word_count = 300

    sliced_ast = slice_ast_for_demo(ast_copy, is_demo=True, max_words=4500)

    total_words = sum(c.word_count or 0 for c in sliced_ast.chapters)
    assert total_words <= 4500
    assert len(sliced_ast.chapters) <= 15

    # Check for callout notice in the last chapter
    last_ch = sliced_ast.chapters[-1]
    has_callout = any(
        isinstance(b, CalloutBlock) and "Demo Preview" in (getattr(b, "title", "") or "")
        for b in last_ch.content
    )
    assert has_callout is True


def test_stage2_ast_slicing_pro_tier_bypasses(twenty_five_page_ast):
    """Verify Pro tier leaves AST 100% untouched."""
    ast_copy = copy.deepcopy(twenty_five_page_ast)
    result = slice_ast_for_demo(ast_copy, is_demo=False)
    assert len(result.chapters) == 25


# ── Stage 3 Tests ─────────────────────────────────────────────────────────────

def test_stage3_output_pdf_capping_and_upsell_page(temp_dir):
    """Verify PyMuPDF strictly caps PDF to 15 pages and appends Page 16 upsell card."""
    compiled_pdf = temp_dir / "output_22_pages.pdf"
    create_test_pdf(compiled_pdf, num_pages=22)

    result: OutputCappingResult = enforce_output_pdf_limit(
        pdf_path=str(compiled_pdf),
        is_demo=True,
        book_title="The Winds of Mystery",
        author_name="Arthur Penhaligon",
        max_pages=15,
    )

    assert result.is_capped is True
    assert result.original_page_count == 22
    assert result.final_page_count == 16  # 15 content pages + 1 upsell page
    assert result.upsell_page_appended is True

    # Re-open capped PDF with PyMuPDF to physically verify structure
    doc = fitz.open(str(compiled_pdf))
    assert len(doc) == 16

    # Verify Page 16 contains upsell text
    last_page = doc[15]
    page_text = last_page.get_text()
    assert "BOOKCRAFT AI — PRO PREVIEW" in page_text
    assert "End of Free 15-Page Preview" in page_text
    assert "The Winds of Mystery" in page_text
    assert "BookCraft Pro" in page_text
    doc.close()


def test_stage3_output_pdf_under_limit_not_capped(temp_dir):
    """Verify that a 12-page PDF is not modified by Stage 3."""
    short_pdf = temp_dir / "output_12_pages.pdf"
    create_test_pdf(short_pdf, num_pages=12)

    result = enforce_output_pdf_limit(
        pdf_path=str(short_pdf),
        is_demo=True,
        max_pages=15,
    )

    assert result.is_capped is False
    assert result.original_page_count == 12
    assert result.final_page_count == 12

    doc = fitz.open(str(short_pdf))
    assert len(doc) == 12
    doc.close()
